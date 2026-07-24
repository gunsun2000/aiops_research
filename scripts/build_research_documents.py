from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "deliverables"
ARCHITECTURE_IMAGE = ROOT / "docs" / "assets" / "architecture_overview.png"

NAVY = "102A43"
BLUE = "174A7E"
LIGHT_BLUE = "E8F0F8"
PALE_BLUE = "F4F7FA"
INK = "17212B"
MUTED = "5D6B78"
LINE = "C9D4DF"
WHITE = "FFFFFF"
POSITIVE = "1F6B4F"
CAUTION = "8A6116"
RISK = "9A2A2A"
FONT = "Malgun Gothic"
MONO_FONT = "Consolas"
PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


@dataclass(frozen=True)
class DocSpec:
    filename: str
    title: str
    subtitle: str
    document_type: str
    short_title: str


REPORT = DocSpec(
    filename="AIOps_4Agent_Research_Report.docx",
    title="AIOps 4-Agent Kubernetes 장애 복구 연구 보고서",
    subtitle="Safety-Bounded Closed-Loop 4-Agent AIOps Framework",
    document_type="대학원 연구 설계 및 실험 결과",
    short_title="AIOps 4-Agent Research Report",
)
OPERATIONS_GUIDE = DocSpec(
    filename="AIOps_Experiment_Operations_Guide.docx",
    title="AIOps 실험 실행 및 검증 가이드",
    subtitle="설치, Control Plane, Mock, Dry-run, Real 실험 재현 절차",
    document_type="연구 재현 및 운영 가이드",
    short_title="AIOps Experiment Operations Guide",
)
POLICY_SPEC = DocSpec(
    filename="AIOps_Agent_Policy_Specification.docx",
    title="4-Agent Action 및 Reward 정책 명세서",
    subtitle="역할, 의사결정, 합의, 보상 및 안전 경계",
    document_type="Agent 정책 설계 문서",
    short_title="AIOps Agent Policy Specification",
)


def _set_run_font(
    run,
    *,
    name: str = FONT,
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr()
    fonts = run._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_style_font(style, *, size: float, color: str, bold: bool) -> None:
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style._element.get_or_add_rPr()
    fonts = style._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)


def _set_paragraph_border(paragraph, *, bottom: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    if bottom:
        node = OxmlElement("w:bottom")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "10")
        node.set(qn("w:space"), "8")
        node.set(qn("w:color"), bottom)
        p_bdr.append(node)


def _set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)


def _set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != PAGE_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {PAGE_WIDTH_DXA}: {widths_dxa}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths_dxa[index])
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _fresh_decimal_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    style_num_id = doc.styles["List Number"].element.pPr.numPr.numId.val
    abstract_num_id = numbering.num_having_numId(style_num_id).abstractNumId.val
    num = numbering.add_num(abstract_num_id)
    num.add_lvlOverride(ilvl=0).add_startOverride(1)
    return num.numId


def _apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = num_id


def _add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = paragraph.add_run("Page ")
    _set_run_font(label, size=8.5, color=MUTED)

    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)

    instruction_run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    instruction_run.append(instruction)

    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)

    value_run = OxmlElement("w:r")
    value = OxmlElement("w:t")
    value.text = "1"
    value_run.append(value)

    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)

    paragraph._p.extend(
        (begin_run, instruction_run, separate_run, value_run, end_run)
    )


def _configure_document(doc: Document, spec: DocSpec, *, compact: bool = False) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    _set_style_font(normal, size=10.3 if compact else 10.5, color=INK, bold=False)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5 if compact else 6)
    normal.paragraph_format.line_spacing = 1.12 if compact else 1.15

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 9, 4),
    ):
        style = doc.styles[name]
        _set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        _set_style_font(style, size=10.3 if compact else 10.5, color=INK, bold=False)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.12 if compact else 1.15

    if "Code Block" not in [style.name for style in doc.styles]:
        code_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = doc.styles["Code Block"]
    _set_style_font(code_style, size=8.6, color="243447", bold=False)
    code_style.font.name = MONO_FONT
    code_style._element.rPr.rFonts.set(qn("w:ascii"), MONO_FONT)
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), MONO_FONT)
    code_style.paragraph_format.left_indent = Inches(0.16)
    code_style.paragraph_format.right_indent = Inches(0.16)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(7)
    code_style.paragraph_format.line_spacing = 1.0
    code_style.paragraph_format.keep_together = True

    header = section.header.paragraphs[0]
    header.text = ""
    header.paragraph_format.space_after = Pt(2)
    header_run = header.add_run(spec.short_title)
    _set_run_font(header_run, size=8.5, bold=True, color=MUTED)
    _set_paragraph_border(header, bottom=LINE)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    _add_page_field(footer)

    first_footer = section.first_page_footer.paragraphs[0]
    first_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(first_footer.add_run("Kyung Hee University | AI-Ops Research"), size=8.5, color=MUTED)

    doc.core_properties.title = spec.title
    doc.core_properties.subject = spec.subtitle
    doc.core_properties.author = "Kyung Hee University AI-Ops Research"
    doc.core_properties.keywords = "AIOps, Multi-Agent, Kubernetes, Chaos Mesh, Prometheus"


def _add_cover(doc: Document, spec: DocSpec) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(44)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(12)
    _set_run_font(kicker.add_run(spec.document_type), size=10.5, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    _set_run_font(title.add_run(spec.title), size=24, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    _set_run_font(subtitle.add_run(spec.subtitle), size=12.5, color=MUTED)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(42)
    _set_paragraph_border(rule, bottom=BLUE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    _set_run_font(meta.add_run("경희대학교 인공지능학과"), size=11, bold=True, color=INK)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta2.paragraph_format.space_after = Pt(4)
    _set_run_font(meta2.add_run("AI Agent 및 AIOps 석사 연구"), size=10.5, color=MUTED)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(date.add_run("2026년 7월"), size=10.5, color=MUTED)
    doc.add_page_break()


def _add_contents(doc: Document, items: Sequence[str]) -> None:
    heading = doc.add_heading("문서 구성", level=1)
    heading.paragraph_format.space_before = Pt(0)
    intro = doc.add_paragraph(
        "이 문서는 다음 순서로 읽도록 구성했습니다. 세부 실행 명령과 원본 설정은 문서 말미의 관련 파일에서 확인할 수 있습니다."
    )
    intro.paragraph_format.space_after = Pt(10)
    _add_numbered_steps(doc, items)
    doc.add_page_break()


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_paragraph(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        _set_run_font(lead, bold=True, color=NAVY)
        run = paragraph.add_run(text[len(bold_lead) :])
        _set_run_font(run)
    else:
        run = paragraph.add_run(text)
        _set_run_font(run)


def _add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        _set_run_font(paragraph.add_run(item))


def _add_numbered_steps(doc: Document, items: Iterable[str]) -> None:
    num_id = _fresh_decimal_numbering(doc)
    for item in items:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15
        _apply_numbering(paragraph, num_id)
        _set_run_font(paragraph.add_run(item))


def _add_callout(doc: Document, title: str, text: str, *, tone: str = "info") -> None:
    palette = {
        "info": (LIGHT_BLUE, BLUE),
        "positive": ("EAF5F0", POSITIVE),
        "caution": ("FFF6DF", CAUTION),
        "risk": ("FCEDED", RISK),
    }
    fill, accent = palette[tone]
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)
    _set_paragraph_shading(paragraph, fill)
    title_run = paragraph.add_run(f"{title}\n")
    _set_run_font(title_run, size=10.2, bold=True, color=accent)
    body_run = paragraph.add_run(text)
    _set_run_font(body_run, size=9.8, color=INK)


def _add_code(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph(style="Code Block")
    _set_paragraph_shading(paragraph, "F1F4F7")
    for index, line in enumerate(code.strip().splitlines()):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        _set_run_font(run, name=MONO_FONT, size=8.6, color="243447")


def _add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths_dxa: Sequence[int],
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_table_geometry(table, widths_dxa)
    _set_repeat_table_header(table.rows[0])

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        _set_run_font(run, size=9.1, bold=True, color=NAVY)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            _set_run_font(run, size=8.9, color=INK)

    _set_table_geometry(table, widths_dxa)
    for row in table.rows:
        _set_row_cant_split(row)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def _add_figure(doc: Document, image_path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    picture = run.add_picture(str(image_path), width=Inches(6.25))
    picture._inline.docPr.set("title", "AIOps 4-Agent research architecture")
    picture._inline.docPr.set("descr", caption)
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(10)
    _set_run_font(caption_paragraph.add_run(caption), size=8.8, color=MUTED, italic=True)


def _add_source_note(doc: Document, paths: Sequence[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    _set_run_font(paragraph.add_run("관련 원본: "), size=8.5, bold=True, color=MUTED)
    _set_run_font(paragraph.add_run(", ".join(paths)), name=MONO_FONT, size=8.2, color=MUTED)


def _build_report() -> Path:
    doc = Document()
    _configure_document(doc, REPORT)
    _add_cover(doc, REPORT)
    _add_contents(
        doc,
        [
            "연구 배경과 문제 정의",
            "연구 목표와 범위",
            "전체 시스템 아키텍처",
            "4-Agent 역할 및 의사결정 구조",
            "안전 경계와 실행 파이프라인",
            "실험 설계",
            "실험 결과 및 해석",
            "연구 기여, 한계, 후속 연구",
        ],
    )

    _add_heading(doc, "1. 연구 배경과 문제 정의")
    _add_paragraph(
        doc,
        "Kubernetes 기반 서비스 운영에서는 장애 탐지, 복구 필요성 판단, 애플리케이션 제어, 인프라 수용성, 비용 제약을 동시에 고려해야 한다. 하나의 Agent가 이 모든 목표를 단독으로 처리하면 판단 근거가 섞이고, 자유 텍스트로 생성된 명령이 곧바로 실행되는 안전 문제도 발생할 수 있다.",
    )
    _add_callout(
        doc,
        "핵심 연구 문제",
        "역할이 분리된 4개 Agent가 동일한 장애 근거를 독립적인 관점에서 검토하고, 합의와 이중 안전 검증을 통과한 bounded action만 Kubernetes에 적용할 수 있는가?",
    )

    _add_heading(doc, "2. 연구 목표와 범위")
    _add_bullets(
        doc,
        [
            "HA, 응용관리, 인프라 운용, 비용 최적화 관점으로 운영 판단을 분리한다.",
            "Agent별 action, reward, 승인·거부 근거를 구조화하여 설명 가능성을 확보한다.",
            "Python Validator와 선택적 Go Guard를 통해 allowlist, replica 제한, 명령 문법을 검증한다.",
            "Chaos Mesh 실제 장애, Prometheus metric, Kubernetes 상태를 이용해 복구 action을 비교한다.",
            "실험 원본을 JSONL로 저장하고 CSV, Markdown, PNG, SVG로 정량 분석한다.",
        ],
    )
    _add_table(
        doc,
        ["범위", "현재 구현", "정확한 경계"],
        [
            ["4-Agent 판단", "구현 완료", "Deterministic 경로와 AutoGen structured 경로를 분리"],
            ["Autonomous evidence", "Mock/test 완료", "FakeEvidenceProvider 중심"],
            ["Kubernetes evidence", "제한적 구현", "Deployment/Pod snapshot 중심"],
            ["Real 장애 제어", "실험 완료", "CLI 확인 절차를 거친 deterministic runner"],
            ["Full evidence fusion", "후속 연구", "Prometheus metric, log enrichment, real-cluster fusion"],
        ],
        [1800, 1860, 5700],
    )
    _add_callout(
        doc,
        "과장 방지 원칙",
        "UI의 4-Agent 판단은 안전한 mock 시연이다. 36회 실제 장애 비교는 deterministic Recovery Runner가 수행했다. AutoGen이 실제 장애 36회를 직접 제어했다고 표현하지 않는다.",
        tone="caution",
    )

    _add_heading(doc, "3. 전체 시스템 아키텍처")
    _add_paragraph(
        doc,
        "시스템은 인프라·관측, Agent 판단, 실행·피드백의 세 계층으로 구성된다. 장애가 발생하면 상태와 metric이 Coordinator로 전달되고, 4개 Agent의 교차 검증 결과가 안전 검증 계층을 거쳐 bounded Kubernetes action으로 변환된다.",
    )
    _add_figure(doc, ARCHITECTURE_IMAGE, "그림 1. AIOps 4-Agent Kubernetes 장애 복구 연구 전체 흐름")
    _add_numbered_steps(
        doc,
        [
            "Chaos Mesh 또는 AIOpsLab이 장애·문제를 제공한다.",
            "Prometheus와 Kubernetes snapshot이 metric, event, deployment/pod 상태를 수집한다.",
            "AI-MCMP Coordinator가 컨텍스트를 4개 전문 Agent에 전달한다.",
            "Agent들은 action과 reward를 제출하고 Coordinator가 합의 결과를 구성한다.",
            "Python Validator와 선택적 Go Guard가 명령과 정책 경계를 검증한다.",
            "Kubernetes 실행 결과를 다시 수집해 복구 성공, 시간, 비용, reward를 평가한다.",
        ],
    )

    doc.add_page_break()
    _add_heading(doc, "4. 4-Agent 역할 및 의사결정 구조")
    _add_table(
        doc,
        ["Agent", "주요 입력", "핵심 판단", "대표 출력"],
        [
            ["HA 지원", "metric, availability, restart", "장애 원인과 복구 필요성", "recovery required / observe"],
            ["응용관리", "HA 판단, service 상태", "bounded 복구 action 제안", "observe / restart / scale"],
            ["인프라 운용", "replica, deployment, capacity", "수용성·배포 안전성", "approve / reject"],
            ["비용 최적화", "action 비용, replica 증가", "과잉 action과 비용 정책", "approve / reject"],
        ],
        [1680, 2300, 2800, 2580],
    )
    _add_paragraph(
        doc,
        "Coordinator는 하나의 거대한 Agent가 아니라 전체 판단 순서와 컨텍스트 전달을 관리하는 오케스트레이터다. 전문 Agent의 역할을 대체하지 않으며, 합의가 실패하면 실행을 중단하거나 더 보수적인 후보로 이동한다.",
    )
    _add_source_note(doc, ["config/agent_registry.json", "src/aiops_k8s_agents/coordinator.py"])

    _add_heading(doc, "5. 안전 경계와 실행 파이프라인")
    _add_table(
        doc,
        ["단계", "검증 항목", "실패 시 처리"],
        [
            ["구조화 action", "허용 action space", "unknown/free-text action 거부"],
            ["Python Validator", "namespace/deployment allowlist, K8s 이름, replica 범위", "실행 차단"],
            ["Go Guard (선택)", "독립 whitelist와 위험 명령 패턴", "최종 실행 차단"],
            ["mock", "명령 생성·검증", "Kubernetes 변경 없음"],
            ["dry-run", "Kubernetes API 호환성", "server dry-run 오류 반환"],
            ["real", "검증된 kubectl action", "stdout/stderr와 상태 snapshot 저장"],
        ],
        [1480, 4800, 3080],
    )
    _add_callout(
        doc,
        "안전 설계의 핵심",
        "LLM 또는 Agent의 자유 텍스트를 셸에 직접 전달하지 않는다. 구조화된 action을 허용된 명령 템플릿으로 변환한 뒤 검증 계층을 통과한 결과만 실행한다.",
        tone="positive",
    )

    _add_heading(doc, "6. 실험 설계")
    _add_heading(doc, "6.1 AIOpsLab 탐지 Benchmark", level=2)
    _add_paragraph(
        doc,
        "Hotel Reservation detection 문제를 대상으로 외부 AIOps benchmark의 문제 시작, metric 수집, 탐지 제출, 평가 결과 저장 경로를 검증했다. 이 실험은 장애 탐지·분석 성능을 확인하며, Kubernetes 복구 action 비교와는 별도 실험이다.",
    )
    _add_table(
        doc,
        ["지표", "결과", "해석"],
        [
            ["전체 실행", "12회", "반복 탐지 실행"],
            ["Correct detection", "12회", "탐지 제출 성공"],
            ["Metric 수집 성공", "11회", "1회는 metric 수집 경계 사례"],
            ["평균 TTD", "4.117초", "탐지 완료까지 평균 시간"],
        ],
        [2500, 1700, 5160],
    )

    _add_heading(doc, "6.2 Chaos Mesh 복구 Action 비교", level=2)
    _add_table(
        doc,
        ["장애", "대상", "관측 metric", "후보 action"],
        [
            ["pod-kill", "paymentservice", "available replica / Pod UID", "observe, restart, scale"],
            ["cpu-stress", "paymentservice", "container CPU rate", "observe, restart, scale"],
            ["memory-stress", "checkoutservice", "working set bytes", "observe, restart, scale"],
            ["network-delay", "paymentservice", "blackbox probe duration", "observe, restart, scale"],
        ],
        [1600, 1900, 2860, 3000],
    )
    _add_callout(
        doc,
        "실험 행렬",
        "장애 4종 x action 3종 x 반복 3회 = 총 36 treatments. 각 treatment는 장애 적용, fault metric 관측, action 실행, 회복 관측, cleanup, JSONL 저장을 독립적으로 수행한다.",
    )

    doc.add_page_break()
    _add_heading(doc, "7. 실험 결과 및 해석")
    _add_table(
        doc,
        ["정책", "CPU", "Memory", "Network", "Pod kill"],
        [
            ["Balanced", "Observe", "Restart", "Restart", "Observe"],
            ["HA first", "Observe", "Restart", "Restart", "Observe"],
            ["Cost first", "Observe", "Observe", "Observe", "Observe"],
            ["Infra first", "Observe", "Restart", "Restart", "Observe"],
        ],
        [1800, 1890, 1890, 1890, 1890],
    )
    _add_bullets(
        doc,
        [
            "Pod kill은 Kubernetes Controller가 자동으로 Pod를 교체하므로 추가 action이 과잉 대응일 수 있었다.",
            "현재 강도와 지속시간의 CPU stress는 자동 종료되어 observe_only가 비용 대비 높은 점수를 받았다.",
            "Memory stress와 network delay에서는 Balanced, HA-first, Infra-first 정책이 rollout_restart를 선택했다.",
            "Cost-first 정책은 자원 변경 비용을 강하게 반영해 모든 장애에서 observe_only를 선택했다.",
            "동일한 실측 결과에서도 reward 가중치에 따라 action ranking이 달라져 reward가 실제 선택에 관여함을 확인했다.",
        ],
    )
    _add_callout(
        doc,
        "검증된 결과",
        "2026년 6월의 36회 본 실험은 36개의 유효 측정과 36개의 복구 성공을 기록했다. 결과는 deterministic Recovery Runner 기반이며, JSONL·CSV·Markdown·그래프로 재현 가능하게 저장된다.",
        tone="positive",
    )

    _add_heading(doc, "8. 연구 기여, 한계, 후속 연구")
    _add_heading(doc, "8.1 연구 기여", level=2)
    _add_bullets(
        doc,
        [
            "역할 기반 4-Agent 운영 판단 구조와 Coordinator 기반 합의 흐름을 구현했다.",
            "구조화 action, allowlist, replica limit, 선택적 이중 Guard로 실행 안전 경계를 만들었다.",
            "AIOpsLab 탐지와 Chaos Mesh real 복구를 분리해 탐지 성능과 제어 효과를 각각 평가했다.",
            "실험 결과를 정책별 reward ranking과 action별 정량 지표로 연결했다.",
        ],
    )
    _add_heading(doc, "8.2 현재 한계", level=2)
    _add_bullets(
        doc,
        [
            "36회 실험의 반복 수가 action별 3회이므로 신뢰구간과 더 큰 표본이 필요하다.",
            "비용은 실제 클라우드 청구액이 아니라 replica와 command 부담을 반영한 정책 점수다.",
            "AutoGen GroupChat의 multi-round real action 선택 비교는 아직 수행하지 않았다.",
            "Single-Agent baseline과 Agent 제거 ablation이 없어 4-Agent의 상대적 우수성을 확정하지 않았다.",
            "KubernetesEvidenceProvider는 deployment/pod snapshot 중심이며 full evidence fusion은 후속 범위다.",
        ],
    )
    _add_heading(doc, "8.3 후속 연구", level=2)
    _add_bullets(
        doc,
        [
            "Single-Agent 및 규칙 기반 baseline과의 탐지 정확도, MTTR, 비용 비교",
            "HA·응용·인프라·비용 Agent 제거 ablation과 정책 민감도 분석",
            "AutoGen multi-round 토론을 실제 장애 action 선택에 연결한 통제 실험",
            "Prometheus metric, Kubernetes event, log enrichment를 결합한 evidence fusion",
            "실제 클라우드 비용, SLO 위반 시간, resource efficiency를 포함한 목적함수 확장",
        ],
    )
    _add_source_note(
        doc,
        [
            "docs/experiments/recovery_action_experiment_guide.md",
            "docs/archive/first_stage_research_completion.md",
        ],
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / REPORT.filename
    doc.save(path)
    return path


def _build_operations_guide() -> Path:
    doc = Document()
    _configure_document(doc, OPERATIONS_GUIDE, compact=True)
    _add_cover(doc, OPERATIONS_GUIDE)
    _add_contents(
        doc,
        [
            "실행 전 원칙과 환경 구분",
            "설치 및 코드 검증",
            "Control Plane 실행",
            "Mock·Dry-run 검증",
            "Prometheus 연결",
            "36회 Real 복구 실험",
            "정량 분석과 결과 확인",
            "AIOpsLab benchmark",
            "문제 해결과 증거 보존",
        ],
    )

    _add_heading(doc, "1. 실행 전 원칙과 환경 구분")
    _add_table(
        doc,
        ["환경", "용도", "사용 시점"],
        [
            ["base", "Conda 기본 환경", "환경 관리만 수행"],
            ["aiops_research", "본 프로젝트 코드·CLI·Recovery 실험", "대부분의 개발과 시험"],
            ["aiopslab", "외부 AIOpsLab 공식 코드", "AIOpsLab benchmark만 실행"],
        ],
        [1800, 4200, 3360],
    )
    _add_callout(
        doc,
        "Real 실행 주의",
        "실제 Kubernetes 변경은 UI에서 수행하지 않는다. CLI의 명시적 mode와 allowlist를 확인한 뒤 개인용 kind cluster에서만 실행한다.",
        tone="risk",
    )

    _add_heading(doc, "2. 설치 및 코드 검증")
    _add_heading(doc, "2.1 최신 코드와 Python 환경", level=2)
    _add_code(
        doc,
        """
cd ~/geonhae/aiops_research
git pull origin master
conda activate aiops_research
python -m pip install -e ".[dev,autogen,ui,docs]"
python -m pytest
""",
    )
    _add_heading(doc, "2.2 Kubernetes context", level=2)
    _add_code(
        doc,
        """
export PATH="$HOME/bin:$PATH"
export KUBECONFIG="$HOME/geonhae/kubeconfigs/kind-geonhae-aiops.yaml"

kubectl config current-context
kubectl get nodes
kubectl get pods -n online-boutique
kubectl get pods -n monitoring-full
""",
    )
    _add_paragraph(doc, "정상 기준: context는 kind-geonhae-aiops이고 node와 주요 Pod가 Ready 상태여야 한다.")
    _add_heading(doc, "2.3 선택적 Go Guard", level=2)
    _add_code(
        doc,
        """
cd ~/geonhae/aiops_research/go/aiops-guard
go test ./...
cd ~/geonhae/aiops_research
""",
    )

    _add_heading(doc, "3. Control Plane 실행")
    _add_code(
        doc,
        """
cd ~/geonhae/aiops_research
conda activate aiops_research

export AIOPS_REPO_ROOT="$(pwd)"
export AIOPS_BIND_ADDRESS="127.0.0.1"
export PORT=18080

aiops-control-plane
""",
    )
    _add_paragraph(doc, "브라우저 주소: http://127.0.0.1:18080/")
    _add_table(
        doc,
        ["화면", "용도"],
        [
            ["대시보드", "전체 연구 흐름과 현재 artifact 상태"],
            ["장애 실험", "장애 4종과 action 3종 실험 설계"],
            ["4-Agent 판단", "안전한 mock 기반 역할별 decision과 consensus"],
            ["안전 검증", "Validator, Guard, allowlist, replica 제한"],
            ["실험 결과", "최근 JSONL, reward ranking, 통계 artifact"],
            ["연구 문서", "DOCX 공식 산출물과 MD 원본"],
        ],
        [2200, 7160],
    )

    _add_heading(doc, "4. Mock·Dry-run 검증")
    _add_heading(doc, "4.1 Autonomous mock", level=2)
    _add_code(
        doc,
        """
aiops-k8s-agents autonomous-run \
  --mode mock \
  --namespace online-boutique \
  --deployment paymentservice \
  --metric cpu \
  --threshold 80 \
  --evidence-value 95 \
  --allowed-namespace online-boutique \
  --allowed-deployment paymentservice
""",
    )
    _add_paragraph(doc, "성공 기준: valid=true이고 실제 Kubernetes resource는 변경되지 않는다.")
    _add_heading(doc, "4.2 Dry-run", level=2)
    _add_paragraph(
        doc,
        "dry-run은 Kubernetes API가 명령을 받아들일 수 있는지 확인하지만 실제 resource를 변경하지 않는다. mock으로 구조를 확인한 뒤 dry-run, 마지막에 real 순서로 진행한다.",
    )

    _add_heading(doc, "5. Prometheus 연결")
    _add_paragraph(doc, "터미널 A에서 포트포워딩을 실행하고 계속 켜 둔다.")
    _add_code(
        doc,
        """
kubectl port-forward \
  -n monitoring-full \
  service/kube-prometheus-stack-prometheus \
  9091:9090
""",
    )
    _add_paragraph(doc, "터미널 B에서 준비 상태와 query를 확인한다.")
    _add_code(
        doc,
        """
export PROM=http://127.0.0.1:9091
curl -sS "$PROM/-/ready"
curl -sSG "$PROM/api/v1/query" --data-urlencode 'query=up'

curl -sSG "$PROM/api/v1/query" \
  --data-urlencode 'query=max(probe_duration_seconds{target="paymentservice"})'

export NETWORK_LATENCY_QUERY='max(probe_duration_seconds{target="paymentservice"})'
""",
    )
    _add_callout(
        doc,
        "연결이 끊긴 경우",
        "Prometheus Pod 재생성으로 port-forward가 끊길 수 있다. 새 Pod가 Ready인지 확인한 뒤 터미널 A에서 port-forward를 다시 시작한다.",
        tone="caution",
    )

    _add_heading(doc, "6. 36회 Real 복구 실험")
    _add_code(
        doc,
        """
cd ~/geonhae/aiops_research

GUARD_BACKEND=go \
MODE=real \
REPETITIONS=3 \
PROMETHEUS_URL="$PROM" \
NETWORK_LATENCY_QUERY="$NETWORK_LATENCY_QUERY" \
bash scripts/server_recovery_action_pilot.sh
""",
    )
    _add_paragraph(
        doc,
        "실행 행렬은 pod-kill, cpu-stress, memory-stress, network-delay 4종과 observe_only, rollout_restart, scale_out 3종을 각각 3회 반복한 총 36 treatments다.",
    )
    _add_paragraph(doc, "장시간 단계에서는 progress wrapper가 elapsed time과 still working 상태를 표시한다.")

    _add_heading(doc, "7. 정량 분석과 결과 확인")
    _add_code(
        doc,
        """
LATEST=$(ls -dt runs/recovery-action-pilot/*/ | head -1)

echo "$LATEST"
wc -l "$LATEST/outcomes.jsonl"
cat "$LATEST/analysis/reward_policy_comparison.md"

bash scripts/server_recovery_statistics.sh
ls "$LATEST/statistics"
cat "$LATEST/statistics/quantitative_summary.md"
""",
    )
    _add_table(
        doc,
        ["산출물", "의미"],
        [
            ["outcomes.jsonl", "Treatment별 원본 metric, action, 성공, 시간, 오류"],
            ["scenario_action_statistics.csv", "장애/action별 평균 복구 시간과 성공률"],
            ["policy_reward_statistics.csv", "reward 정책별 action ranking"],
            ["mean_recovery_seconds_by_action.png", "평균 복구 시간 시각화"],
            ["success_rate_by_action.png", "복구 성공률 시각화"],
            ["reward_by_policy.png", "정책별 reward 점수 시각화"],
        ],
        [3600, 5760],
    )
    _add_code(
        doc,
        """
python - "$LATEST/outcomes.jsonl" <<'PY'
import json, sys
for line in open(sys.argv[1], encoding="utf-8"):
    row = json.loads(line)
    if not row.get("measurement_valid"):
        print(row.get("treatment_id"), row.get("error"))
PY
""",
    )
    _add_paragraph(doc, "위 명령에서 아무것도 출력되지 않으면 모든 measurement가 유효하다.")

    _add_heading(doc, "8. AIOpsLab benchmark")
    _add_code(
        doc,
        """
cd ~/geonhae/aiops_research
conda activate aiopslab

bash scripts/server_aiopslab_auto_detection.sh
bash scripts/server_aiopslab_repeat_detection.sh
bash scripts/server_aiopslab_summarize_runs.sh
""",
    )
    _add_callout(
        doc,
        "실험 목적 분리",
        "AIOpsLab은 장애 탐지·분석 benchmark이고, Chaos Mesh recovery matrix는 Kubernetes 복구 action 비교다. 결과를 한 실험처럼 합치지 않는다.",
    )

    _add_heading(doc, "9. 문제 해결과 증거 보존")
    _add_table(
        doc,
        ["증상", "확인", "조치"],
        [
            ["kubectl 권한 오류", "current-context와 KUBECONFIG", "개인 kind kubeconfig 재설정"],
            ["Prometheus 연결 거부", "Pod Ready와 port-forward", "포트포워딩 재시작"],
            ["query sample 없음", "metric name과 target label", "실제 label로 query 수정"],
            ["36줄 미만", "outcomes.jsonl와 stderr", "실패 treatment만 재실행"],
            ["그래프 없음", "statistics 폴더", "server_recovery_statistics.sh 실행"],
        ],
        [2300, 3300, 3760],
    )
    _add_bullets(
        doc,
        [
            "실험 시작 전 Git commit, kube context, Pod 상태를 기록한다.",
            "실험 중 생성된 JSONL을 원본 evidence로 보존한다.",
            "실패 결과를 삭제하지 말고 measurement_valid와 error를 함께 남긴다.",
            "발표에는 요약 그래프를 사용하되 원본 경로를 함께 기록한다.",
        ],
    )
    _add_source_note(
        doc,
        [
            "docs/submission/execution_code_guide.md",
            "docs/submission/test_guide.md",
            "docs/submission/control_plane_ui_guide.md",
        ],
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / OPERATIONS_GUIDE.filename
    doc.save(path)
    return path


def _build_policy_spec() -> Path:
    doc = Document()
    _configure_document(doc, POLICY_SPEC, compact=True)
    _add_cover(doc, POLICY_SPEC)
    _add_contents(
        doc,
        [
            "정책 목적과 개념",
            "Agent Registry",
            "공통 입출력 계약",
            "Agent별 판단 정책",
            "Consensus와 action 선택",
            "Reward 정책",
            "안전 검증 정책",
            "장애별 action 해석",
            "정책 한계와 확장",
        ],
    )

    _add_heading(doc, "1. 정책 목적과 개념")
    _add_paragraph(
        doc,
        "본 명세는 4-Agent가 동일한 장애 evidence를 역할별로 검토하고, 허용된 Kubernetes recovery action을 선택하는 기준을 정의한다. Reward는 현재 강화학습용 학습 보상이 아니라 action 후보를 비교하는 정책 평가 점수다.",
    )
    _add_table(
        doc,
        ["용어", "정의"],
        [
            ["Evidence", "metric, threshold, deployment/pod snapshot, event와 실행 결과"],
            ["Cause", "cpu_saturation, memory_saturation, network_degradation 등 구조화 원인"],
            ["Action", "observe_only, rollout_restart, scale_out의 bounded 후보"],
            ["Decision", "Agent별 action, approved, reward, reason"],
            ["Consensus", "Agent 검토 결과를 결합한 최종 실행 합의"],
            ["Reward", "HA·응용·인프라·비용 관점의 정책 평가 점수"],
        ],
        [2300, 7060],
    )

    _add_heading(doc, "2. Agent Registry")
    _add_table(
        doc,
        ["Agent", "책임", "허용 action 예"],
        [
            ["AIServiceHASupportAgent", "장애 원인, 심각도, 복구 필요성", "ha_recovery_required, ha_no_action"],
            ["AIApplicationManagementAgent", "서비스 복구 action 후보 제안", "app_observe_only, app_rollout_restart, app_scale_deployment"],
            ["AISemiconductorInfraOpsAgent", "replica·deployment·capacity 검증", "infra_capacity_approved, infra_capacity_rejected"],
            ["CostOptimizationAgent", "비용·과잉 action 검토", "cost_budget_approved, cost_budget_rejected"],
        ],
        [2400, 3600, 3360],
    )
    _add_paragraph(
        doc,
        "Registry는 역할 설명만 저장하는 목록이 아니라 Agent가 제출할 수 있는 action boundary를 정의한다. Registry 밖 action은 최종 실행 후보가 될 수 없다.",
    )
    _add_source_note(doc, ["config/agent_registry.json", "src/aiops_k8s_agents/agent_registry.py"])

    _add_heading(doc, "3. 공통 입출력 계약")
    _add_code(
        doc,
        """
AgentDecision {
  agent: string
  action: string
  approved: boolean
  reward: number
  reason: string
  metadata: object
}
""",
    )
    _add_bullets(
        doc,
        [
            "approved=true는 해당 Agent 관점에서 현재 action을 허용한다는 의미다.",
            "approved=false는 실행을 차단하거나 더 보수적인 후보를 선택해야 함을 의미한다.",
            "reason은 metric, threshold, replica, 비용 조건을 포함하는 설명 가능한 근거다.",
            "unknown metric은 cpu_saturation으로 간주하지 않고 unknown_metric으로 처리한다.",
            "unknown_metric의 기본 action은 observe_only이며 즉시 scale_out하지 않는다.",
        ],
    )

    _add_heading(doc, "4. Agent별 판단 정책")
    _add_heading(doc, "4.1 HA 지원 Agent", level=2)
    _add_table(
        doc,
        ["Metric 방향", "장애 조건", "Cause", "기본 판단"],
        [
            ["high_is_bad", "value >= threshold", "CPU/Memory/Latency saturation", "복구 후보 검토"],
            ["low_is_bad", "value < threshold", "low_availability", "복구 후보 검토"],
            ["unknown", "정책 매칭 실패", "unknown_metric", "observe_only"],
        ],
        [1800, 2200, 2780, 2580],
    )
    _add_callout(
        doc,
        "Low-is-bad 처리",
        "availability, available_replicas, ready_replicas는 값이 낮을수록 위험하다. evidence 진단은 metric policy의 signal_direction을 사용한다.",
    )

    _add_heading(doc, "4.2 응용관리 Agent", level=2)
    _add_table(
        doc,
        ["Action", "실행 의미", "적용 조건"],
        [
            ["observe_only", "변경 없이 Kubernetes 자체 복구 관찰", "일시 장애, 자동 회복, unknown metric"],
            ["rollout_restart", "Deployment rollout restart", "memory/latency/network 상태 초기화 후보"],
            ["scale_out", "Replica를 정책 범위 안에서 증가", "지속 부하 또는 가용성 부족"],
        ],
        [2200, 3500, 3660],
    )

    _add_heading(doc, "4.3 인프라 운용 Agent", level=2)
    _add_bullets(
        doc,
        [
            "요청 replica가 configured maximum을 넘는지 검토한다.",
            "namespace와 deployment가 실행 대상 정책에 포함되는지 확인한다.",
            "현재 연구에서는 Kubernetes replica/deployment 안전성이 중심이며 실제 GPU/NPU 스케줄링은 포함하지 않는다.",
        ],
    )

    _add_heading(doc, "4.4 비용 최적화 Agent", level=2)
    _add_bullets(
        doc,
        [
            "Replica 증가량과 실행 command 수를 비용·운영 부담으로 평가한다.",
            "불필요한 scale_out보다 observe_only 또는 restart를 선호할 수 있다.",
            "현재 비용은 실제 클라우드 청구액이 아닌 정책 점수다.",
        ],
    )

    _add_heading(doc, "5. Consensus와 Action 선택")
    _add_numbered_steps(
        doc,
        [
            "HA Agent가 원인과 복구 필요성을 제시한다.",
            "응용관리 Agent가 bounded action 후보를 만든다.",
            "인프라와 비용 Agent가 각 후보의 제약을 검토한다.",
            "Coordinator가 승인·거부와 reward를 모아 실행 가능한 후보를 정렬한다.",
            "Validator가 실패한 후보는 executed_actions에 넣지 않는다.",
            "모든 후보가 실패하면 safe failure와 requires_human_review=true를 반환한다.",
        ],
    )
    _add_callout(
        doc,
        "독립 판단의 의미",
        "Agent가 자유로운 장시간 대화를 하지 않더라도 동일한 evidence에 대해 역할별 판단을 독립적으로 제출하고 Coordinator가 결과를 결합할 수 있다. AutoGen GroupChat은 이 구조의 보조 경로다.",
    )

    _add_heading(doc, "6. Reward 정책")
    _add_table(
        doc,
        ["정책", "HA", "응용", "인프라", "비용"],
        [
            ["balanced", "0.30", "0.30", "0.20", "0.20"],
            ["ha_first", "0.50", "0.25", "0.15", "0.10"],
            ["cost_first", "0.25", "0.20", "0.15", "0.40"],
            ["infra_first", "0.25", "0.20", "0.40", "0.15"],
        ],
        [2400, 1740, 1740, 1740, 1740],
    )
    _add_code(
        doc,
        """
ActionScore
= HA score * HA weight
  + Application score * Application weight
  + Infrastructure score * Infrastructure weight
  + Cost score * Cost weight
""",
    )
    _add_paragraph(
        doc,
        "Predicted reward는 정책 가중치를 적용한 선택 점수이고, observed outcome score는 실제 treatment 결과를 네 Agent 관점에서 재평가한 평균 점수다.",
    )
    _add_table(
        doc,
        ["Agent signal", "예시 Reward", "의미"],
        [
            ["ha_recovery_required", "+0.90", "가용성 관점에서 복구 필요"],
            ["ha_no_action", "+0.20", "불필요한 변경 회피"],
            ["app_rollout_restart", "+0.85", "상태 초기화 action 지지"],
            ["app_scale_deployment", "+0.85", "scale-out action 지지"],
            ["infra_capacity_approved", "+0.70", "인프라 수용 가능"],
            ["infra_capacity_rejected", "-0.60", "capacity 제약으로 차단"],
            ["cost_budget_approved", "+0.60", "비용 정책 범위"],
            ["cost_budget_rejected", "-0.70", "비용 초과 위험"],
        ],
        [3350, 1650, 4360],
    )

    _add_heading(doc, "7. 안전 검증 정책")
    _add_table(
        doc,
        ["경계", "검증", "거부 예"],
        [
            ["Registry", "Agent별 허용 action", "등록되지 않은 action"],
            ["Allowlist", "namespace/deployment", "공용·시스템 namespace"],
            ["Naming", "Kubernetes resource name", "shell metacharacter 포함"],
            ["Replica", "min/max 범위", "상한 초과 scale_out"],
            ["Command template", "허용 kubectl 문법", "임의 shell command"],
            ["Go Guard (선택)", "독립 whitelist 재검증", "위험 명령 패턴"],
        ],
        [2200, 3500, 3660],
    )
    _add_callout(
        doc,
        "Autonomous safety",
        "Autonomous mode에서도 validation을 통과하지 못한 action은 실행 목록에 포함되지 않으며, 안전한 대안이 없으면 사람 검토를 요구한다.",
        tone="positive",
    )

    _add_heading(doc, "8. 장애별 Action 해석")
    _add_table(
        doc,
        ["장애", "관찰이 유리한 조건", "재시작이 유리한 조건", "Scale-out이 유리한 조건"],
        [
            ["Pod kill", "Controller가 즉시 대체", "상태가 회복되지 않음", "가용 replica 지속 부족"],
            ["CPU stress", "짧은 부하가 자동 종료", "프로세스 stuck", "지속 포화와 처리량 부족"],
            ["Memory stress", "압박이 일시적", "상태 초기화 효과", "동시 요청 분산 필요"],
            ["Network delay", "장애가 자동 종료", "영향 Pod 교체 효과", "Replica 분산이 지연 완화"],
        ],
        [1500, 2620, 2620, 2620],
    )
    _add_paragraph(
        doc,
        "위 표는 고정 정답이 아니라 정책 해석 기준이다. 실제 선택은 metric 회복, 복구 시간, command 수, replica 증가량, safety 결과를 함께 평가한다.",
    )

    _add_heading(doc, "9. 정책 한계와 확장")
    _add_bullets(
        doc,
        [
            "Reward 계수는 초기 정책 기준선이며 더 큰 실험 표본으로 보정해야 한다.",
            "실제 비용, SLO 위반 시간, error budget을 직접 반영하는 목적함수가 필요하다.",
            "Single-Agent baseline과 Agent ablation으로 각 역할의 기여를 분리해야 한다.",
            "AutoGen multi-round 판단을 deterministic runner와 동일 조건에서 비교해야 한다.",
            "정책 파일과 Registry 버전을 실험 결과에 함께 저장해 재현성을 강화해야 한다.",
        ],
    )
    _add_source_note(
        doc,
        [
            "config/agent_decision_policy.json",
            "config/recovery_action_experiments.json",
            "docs/design/agent_action_reward_policy.md",
        ],
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / POLICY_SPEC.filename
    doc.save(path)
    return path


def build_all() -> list[Path]:
    return [_build_report(), _build_operations_guide(), _build_policy_spec()]


if __name__ == "__main__":
    for output in build_all():
        print(output.relative_to(ROOT))
