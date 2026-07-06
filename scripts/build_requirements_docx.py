from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/submission/requirements_definition.docx")


def set_font(
    run,
    name: str = "Malgun Gothic",
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_font(run, size=9.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        set_cell_shading(cell, fill)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True, fill="F2F4F7")
        cell.width = Inches(widths[index])
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            cells[index].width = Inches(widths[index])
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_font(run, size=16 if level == 1 else 13, bold=True, color="2E74B5")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_font(paragraph.add_run(item), size=10.5)


def build_docx() -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(
        title.add_run("AIOps 4-Agent Kubernetes 장애 복구 연구 요구사항 정의서"),
        size=18,
        bold=True,
        color="0B2545",
    )

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(
        subtitle.add_run("Safety-Bounded Closed-Loop Autonomous 4-Agent AIOps Framework"),
        size=11,
        color="555555",
    )

    doc.add_paragraph(
        "본 문서는 Kubernetes 서비스 장애 감시와 복구 자동화를 위한 4-Agent AIOps 연구 프로토타입의 구현 범위, 기능 요구사항, 시험 방법, 산출물 구성을 정리한다."
    )

    add_heading(doc, "1. 연구 목적")
    add_bullets(
        doc,
        [
            "4개의 AI Agent가 역할별로 장애를 판단하고 복구 action을 제안한다.",
            "Action/Reward 정책으로 Agent 판단을 교차 검증한다.",
            "Python Validator와 선택적 Go Guard를 통해 위험한 Kubernetes action을 차단한다.",
            "Chaos Mesh, AIOpsLab, Prometheus, Kubernetes 환경에서 실험 결과를 수집한다.",
        ],
    )

    add_heading(doc, "2. 기능 요구사항")
    add_table(
        doc,
        ["ID", "요구사항", "상태"],
        [
            ["FR-01", "장애 입력을 구조화된 AlertEvent로 변환", "완료"],
            ["FR-02", "HA Agent의 장애 원인 및 복구 필요성 판단", "완료"],
            ["FR-03", "Application Agent의 bounded recovery action 후보 생성", "완료"],
            ["FR-04", "Infrastructure Agent의 replica/deployment 안전성 검토", "완료"],
            ["FR-05", "Cost Agent의 비용 및 과잉 action 검토", "완료"],
            ["FR-06", "Python Validator 기반 allowlist와 replica limit 검증", "완료"],
            ["FR-07", "선택적 Go Guard 기반 이중 안전 검증", "완료"],
            ["FR-08", "mock, dry-run, real 실행 모드 제공", "완료"],
            ["FR-09", "Chaos Mesh 기반 4종 장애 반복 실험", "완료"],
            ["FR-10", "JSONL/CSV/Markdown/PNG/SVG 결과 산출", "완료"],
        ],
        [0.8, 5.0, 0.8],
    )

    add_heading(doc, "3. Agent 역할")
    add_table(
        doc,
        ["Agent", "역할", "대표 action"],
        [
            ["AIServiceHASupportAgent", "서비스 장애 진단, 가용성 판단, 복구 필요성 평가", "ha_scale_out_required"],
            ["AIApplicationManagementAgent", "observe, restart, scale-out action 후보 생성", "app_scale_deployment"],
            ["AISemiconductorInfraOpsAgent", "Kubernetes replica/deployment 안전성 검토", "infra_capacity_approved"],
            ["CostOptimizationAgent", "비용 증가와 과잉 action 검토", "cost_budget_approved"],
        ],
        [2.0, 3.5, 1.2],
    )

    add_heading(doc, "4. 시험 방법")
    add_table(
        doc,
        ["시험", "명령", "성공 기준"],
        [
            ["Python 단위 테스트", "python -m pytest", "전체 테스트 passed"],
            ["Go Guard 테스트", "go test ./...", "Go guard 테스트 통과"],
            ["Agent Registry", "aiops-k8s-agents list-agents", "4개 Agent 조회"],
            ["Autonomous mock", "aiops-k8s-agents autonomous-run", "valid=true"],
            ["Recovery 실험", "server_recovery_action_pilot.sh", "36개 outcome 기록"],
            ["정량 분석", "server_recovery_statistics.sh", "PNG/SVG/CSV/JSON 산출"],
        ],
        [1.5, 2.7, 2.3],
    )

    add_heading(doc, "5. 산출물")
    add_table(
        doc,
        ["산출물", "위치"],
        [
            ["요구사항 정의서", "docs/submission/requirements_definition.md"],
            ["Word 요구사항 정의서", "docs/submission/requirements_definition.docx"],
            ["설치 및 실행 가이드", "docs/submission/install_and_run_guide.md"],
            ["시험 가이드", "docs/submission/test_guide.md"],
            ["Agent Registry", "config/agent_registry.json"],
            ["Action/Reward 정책", "docs/design/agent_action_reward_policy.md"],
            ["Recovery 실험 결과", "runs/recovery-action-pilot/<run>/"],
            ["정량 분석 그래프", "runs/recovery-action-pilot/<run>/statistics/"],
        ],
        [2.4, 4.1],
    )

    add_heading(doc, "6. 향후 확장")
    add_bullets(
        doc,
        [
            "single-agent baseline 비교",
            "Agent 제거 ablation 실험",
            "reward 민감도 분석",
            "AutoGen multi-round real action 선택",
            "Prometheus metric과 log enrichment 기반 full evidence fusion",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("AIOps 4-Agent Research"), size=8, color="777777")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
